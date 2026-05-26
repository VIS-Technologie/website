"""
Generuje brief treściowy dla klienta Vis Technologie.

Wyjście: docs/brief-tresci-vistechnologie.docx
Wymaga:   python-docx (pip install python-docx)
Uruchom:  python tools/generate-brief.py

Treść briefu (pytania, cytaty ze starej strony) zaszyta poniżej w stałych
PART_1, PART_2, EPILOG. Edytuj tam i uruchom ponownie, żeby przegenerować.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ----------------------------------------------------------------------------
# Treść briefu
# ----------------------------------------------------------------------------

PART_1_TITLE = "CZĘŚĆ 1 — Strategia i ton"
PART_1_INTRO = (
    "Decyzje w tej części ustawiają ton całej strony. "
    "Wypełnij ją najpierw — będzie łatwiej odpowiadać "
    "na pytania o konkretne sekcje w Części 2."
)

# Lista sekcji Części 1: (numer, tytuł, lista pytań).
PART_1 = [
    (
        "1.1",
        "Grupa docelowa",
        [
            "Kto jest waszym idealnym klientem? (branża, wielkość firmy, region)",
            "Kto na stronie szuka informacji — decydent biznesowy, czy techniczny (IT, kierownik terminala)?",
            "Czy strona ma też przyciągać kandydatów do pracy lub partnerów?",
        ],
    ),
    (
        "1.2",
        "USP / wyróżnik",
        [
            "Jednym zdaniem: czym Vis Technologie różni się od konkurencji?",
            "3 rzeczy, które najczęściej powtarzacie klientom, żeby ich przekonać.",
        ],
    ),
    (
        "1.3",
        "Ton komunikacji",
        [
            "Formalny / partnerski / techniczny — które z tych trzech jest najbliższe? (jeden wybór)",
            "Przykład „tak, takim językiem mówimy” — wklej fragment z innego materiału (oferta, e-mail, broszura).",
            "Przykład „nie, tak nie chcemy brzmieć” — opcjonalnie.",
        ],
    ),
    (
        "1.4",
        "Języki strony",
        [
            "Polski only, czy też angielski (drugi język)?",
            "Jeśli EN — czy dostarczycie tłumaczenia, czy spodziewacie się ich od nas?",
        ],
    ),
    (
        "1.5",
        "Konkurenci i strony, które lubicie",
        [
            "2–3 linki do stron konkurencyjnych firm.",
            "2–3 linki do stron (dowolnej branży), które wam się estetycznie lub treściowo podobają — i co konkretnie się podoba.",
        ],
    ),
]


PART_2_TITLE = "CZĘŚĆ 2 — Treść sekcji nowej strony"
PART_2_INTRO = (
    "Sekcje w kolejności, w jakiej pojawią się na stronie. "
    "Dla każdej sekcji: po co ona jest — co macie teraz na starej stronie "
    "— pytania i pola do wypełnienia. Limity znaków przy pytaniach są "
    "wskazówkami, nie regułami. Można odpowiadać w punktach."
)

# Lista sekcji Części 2.
# (numer, tytuł, "po co", cytat-ze-starej-lub-None, [pytania])
PART_2 = [
    (
        "2.1",
        "Hero",
        "Pierwsze 5 sekund — kto jesteście i co dajecie.",
        (
            "VIS technologie. Firma powstała w Trójmieście, gdzie mieści się jej główna "
            "siedziba. Swoją działalność opiera o wieloletnie doświadczenie założycieli "
            "i współpracowników w branży informatycznej oraz ich pasję do rozwijającej się "
            "branży jaką jest logistyka intermodalna.\n\n"
            "VIS PYXIS jest systemem informatycznym, opracowanym jako innowacyjne, "
            "zintegrowane narzędzie usprawniające zarządzanie procesami w logistyce "
            "intermodalnej, mające na celu podniesienie wydajności przez automatyzację "
            "procesów…"
        ),
        "vistechnologie.filmove.kylos.pl, slider na górze strony",
        [
            "Główny nagłówek (max 8 słów) — co ma rzucić się w oczy w 1. sekundzie?",
            "Podtytuł (max 25 słów) — co rozwija nagłówek?",
            "Tekst na przycisku CTA + dokąd prowadzi (np. „Zobacz VIS PYXIS” → sekcja Case study, „Napisz do nas” → Kontakt).",
            "Czy chcecie 2 CTA (główne + drugorzędne) czy 1?",
        ],
    ),
    (
        "2.2",
        "Stats (liczby)",
        "Szybki dowód kompetencji — 3–5 liczb z opisem.",
        None,
        None,
        [
            "Wpisz 3–5 metryk, którymi chcecie się chwalić (przykłady: „X lat na rynku”, „Y obsłużonych terminali”, „Z wdrożeń”, „W klientów” — ale tylko prawdziwe).",
            "Dla każdej liczby: wartość + 2–4 słowa opisu pod nią.",
        ],
    ),
    (
        "2.3",
        "Services / Co robimy",
        "Lista usług, które oferujecie poza produktem VIS PYXIS.",
        (
            "Czym się zajmujemy.\n"
            "Aplikacje webowe — Specjalizujemy się w aplikacjach webowych z wykorzystaniem "
            "technologii ASP.Net oraz DevExpress.\n"
            "Aplikacje desktopowe — Specjalizujemy się w aplikacjach desktopowych "
            "z wykorzystaniem technologii WinForms, Windows Presentation Foundation, oraz DevExpress.\n"
            "Aplikacje mobilne — Specjalizujemy się w tworzeniu aplikacji mobilnych na system Android."
        ),
        "vistechnologie.filmove.kylos.pl, sekcja „Czym się zajmujemy”",
        [
            "Czy ta lista usług jest aktualna? (T/N + komentarz)",
            "Co dodać (np. konsulting, integracje, wdrożenia, utrzymanie)?",
            "Co usunąć?",
            "Dla każdej usługi: nagłówek + opis 2–3 zdania + dla kogo to.",
        ],
    ),
    (
        "2.4",
        "Process / How We Work",
        "Pokazać klientowi, jak wygląda współpraca krok po kroku — redukuje niepewność.",
        None,
        None,
        [
            "Opiszcie 3–5 etapów współpracy z klientem od pierwszego kontaktu do wdrożenia.",
            "Dla każdego etapu: krótka nazwa + 1–2 zdania co się dzieje + ile to trwa (orientacyjnie).",
        ],
    ),
    (
        "2.5",
        "Case study: VIS PYXIS",
        "Flagowy produkt — najważniejszy dowód, że umiecie wdrażać.",
        (
            "VIS PYXIS — zintegrowany system do obsługi logistyki intermodalnej.\n\n"
            "Charakterystyka: Wiarygodność (m.in. wiarygodna prognoza sprzedaży, planowanie tras), "
            "Efektywność ekonomiczna (monitorowanie kosztów, generowanie listów przewozowych), "
            "Efektywność techniczna (budowa pociągu, weryfikacja stanu kontenerów).\n\n"
            "Funkcjonalność: Wysoki poziom dokładności obsługi terminali, Zarządzanie zleceniami, "
            "Zarządzanie trasami, Kontrakty, Inwentaryzacja kontenerów, Inwentaryzacja pociągu.\n\n"
            "Moduły: obsługi sprzedaży, terminalowy, finansowy, obsługi klienta, operacyjny, techniczny."
        ),
        "vistechnologie.filmove.kylos.pl, sekcje „Opis Rozwiązania” … „Korzyści”",
        [
            "Krótki opis VIS PYXIS w 2–3 zdaniach (do podtytułu sekcji).",
            "Które z 6 modułów funkcjonalnych zostawiamy? (do odhaczenia: terminale / zlecenia / trasy / kontrakty / inwentaryzacja kontenerów / inwentaryzacja pociągu)",
            "Które z 6 modułów systemowych zostawiamy? (sprzedaży / terminalowy / finansowy / obsługi klienta / operacyjny / techniczny)",
            "Top 3 korzyści, które chcecie wyeksponować na pierwszym ekranie sekcji.",
            "Czy jest screen / wideo / demo produktu, które możemy pokazać?",
            "Czy są klienci, których wdrożenia można wymienić z nazwy?",
        ],
    ),
    (
        "2.6",
        "O nas",
        "Budowa zaufania — kto stoi za firmą.",
        (
            "Firma VIS TECHNOLOGIE powstała w Trójmieście, gdzie mieści się jej główna "
            "siedziba. Swoją działalność opiera o wieloletnie doświadczenie założycieli "
            "i współpracowników w branży informatycznej oraz ich pasję do rozwijającej się "
            "branży jaką logistyka intermodalna…"
        ),
        "vistechnologie.filmove.kylos.pl, sekcja „O nas”",
        [
            "Aktualizacja: kiedy powstała firma (rok), gdzie siedziba (potwierdzenie / zmiana).",
            "Ilu osób zespół? Jakie kompetencje (np. „3 architektów .NET, 2 DevOps, …”)?",
            "Czy chcecie wymienić założycieli z imienia/nazwiska + zdjęcia?",
            "1 zdanie historii — od czego zaczęliście, dokąd idziecie.",
        ],
    ),
    (
        "2.7",
        "Misja",
        "Krótki manifest — czemu robicie to co robicie.",
        (
            "Mając na uwadze stale rosnące oczekiwania naszych Klientów, jak również prężny "
            "rozwój branży logistycznej, stawiamy sobie za cel podążanie za najnowocześniejszymi "
            "technologiami i rozwiązaniami. Celem funkcjonowania firmy jest osiągnięcie na rynku "
            "pozycji partnera strategicznego dla swoich klientów… (pełny tekst kilkukrotnie dłuższy)"
        ),
        "vistechnologie.filmove.kylos.pl, sekcja „Misja firmy”",
        [
            "Czy chcecie skrócić obecną misję do 1–3 zdań? (jeśli tak — propozycja)",
            "Albo czy chcecie zostawić długą wersję na osobnej podstronie?",
        ],
    ),
    (
        "2.8",
        "Klienci / Referencje / Partnerzy",
        "„Social proof” — pokazać, że inni już wam zaufali.",
        None,
        None,
        [
            "Lista klientów do pokazania (nazwa firmy + zgoda na publikację — T/N).",
            "Czy chcecie logotypy klientów na stronie? (jeśli tak — dostarczcie SVG/PNG)",
            "Cytaty referencyjne — macie 1–3 wypowiedzi klientów? (osoba, stanowisko, firma, cytat 2–4 zdania)",
            "Partnerstwa / certyfikaty (Microsoft Partner, ISO, inne) — lista + logotypy.",
        ],
    ),
    (
        "2.9",
        "Aktualności / Blog",
        "Sygnalizuje aktywność firmy + SEO.",
        None,
        None,
        [
            "Czy chcecie sekcję aktualności / bloga na nowej stronie? (T/N)",
            "Jeśli tak: jak często będziecie publikować (raz w miesiącu / kwartale / sporadycznie)?",
            "Kto będzie pisać (wewnętrznie / zewnętrznie / pomoc agencji)?",
            "Jakie tematy chcielibyście poruszać (3–5 przykładów)?",
            "Czy są już teksty / case studies / branżowe komentarze, które można od razu opublikować na start?",
        ],
    ),
    (
        "2.10",
        "Kontakt",
        "Zamknięcie strony — jak się z wami skontaktować.",
        (
            "Klonowa 4/6, 80-297 Banino\n"
            "biuro@vistechnologie.pl\n"
            "tel. 887 362 098"
        ),
        "vistechnologie.filmove.kylos.pl, stopka i sekcja „Kontakt”",
        [
            "Czy adres jest aktualny?",
            "Czy numer telefonu jest aktualny? Czy chcecie kilka numerów (np. sprzedaż / wsparcie)?",
            "Czy chcecie formularz kontaktowy na stronie? Jakie pola (imię, e-mail, firma, temat, treść)?",
            "Gdzie ma trafiać wiadomość z formularza (e-mail / CRM)?",
            "Godziny pracy (do wyświetlenia)?",
            "Social media — LinkedIn, Facebook, inne? (linki)",
        ],
    ),
    (
        "2.11",
        "Cookies / RODO / Prawne",
        "Wymogi prawne — bez tego strona nie idzie na produkcję.",
        None,
        None,
        [
            "Czy macie już dokumenty: Polityka prywatności / Regulamin / Polityka cookies? Jeśli tak — dostarczcie pliki (.pdf / .docx). Jeśli nie — czy mamy pomóc przygotować szablony? (uwaga: nie jest to porada prawna)",
            "Treść bannera cookies (1–2 zdania) — wasza wersja, czy nasza propozycja?",
            "Dane administratora danych osobowych (firma, adres, kontakt) — potwierdzenie tych z sekcji Kontakt czy inne?",
        ],
    ),
]


EPILOG_TITLE = "Dopiski"
EPILOG = [
    "Co jeszcze powinno być na nowej stronie? — pole otwarte na rzeczy, których nie ujęliśmy.",
    "Co kasujemy ze starej strony? — wprost wymieńcie to, czego nie chcecie przenosić.",
    "Termin wypełnienia briefu — do kiedy chcecie nam odesłać?",
]


# ----------------------------------------------------------------------------
# Pomocnicze: formatowanie
# ----------------------------------------------------------------------------

GRAY_BG = "F0F0F0"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex: str = "BFBFBF", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_quote_box(doc: Document, text: str, source: str) -> None:
    """Szary box z cytatem ze starej strony + mała linijka źródła."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, GRAY_BG)
    set_cell_border(cell, "D9D9D9")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # wyczyść domyślny pusty paragraf
    cell.paragraphs[0].text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    source_p = doc.add_paragraph()
    src_run = source_p.add_run(f"Źródło: {source}")
    src_run.italic = True
    src_run.font.size = Pt(9)
    src_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_answer_box(doc: Document, placeholder: str = "[tutaj wpisz odpowiedź]") -> None:
    """Biała ramka z placeholderem dla odpowiedzi klienta."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_border(cell, "808080")
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    run = p.add_run(placeholder)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    # dodaj 2 puste linie żeby ramka miała wysokość
    cell.add_paragraph()
    cell.add_paragraph()


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document) -> None:
    """Pole TOC — Word zaktualizuje przy otwarciu (Ctrl+A, F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_text = OxmlElement("w:t")
    fld_char_text.text = "(Spis treści zostanie wypełniony po otwarciu w Word: Ctrl+A, F9)"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_text)
    run._r.append(fld_char_end)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Strona ")
    fld1_begin = OxmlElement("w:fldChar")
    fld1_begin.set(qn("w:fldCharType"), "begin")
    fld1_instr = OxmlElement("w:instrText")
    fld1_instr.set(qn("xml:space"), "preserve")
    fld1_instr.text = "PAGE"
    fld1_end = OxmlElement("w:fldChar")
    fld1_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld1_begin)
    run._r.append(fld1_instr)
    run._r.append(fld1_end)
    run2 = p.add_run(" z ")
    fld2_begin = OxmlElement("w:fldChar")
    fld2_begin.set(qn("w:fldCharType"), "begin")
    fld2_instr = OxmlElement("w:instrText")
    fld2_instr.set(qn("xml:space"), "preserve")
    fld2_instr.text = "NUMPAGES"
    fld2_end = OxmlElement("w:fldChar")
    fld2_end.set(qn("w:fldCharType"), "end")
    run2._r.append(fld2_begin)
    run2._r.append(fld2_instr)
    run2._r.append(fld2_end)


# ----------------------------------------------------------------------------
# Builder dokumentu
# ----------------------------------------------------------------------------


def build_section(
    doc: Document,
    number: str,
    title: str,
    purpose: str | None,
    quote: str | None,
    quote_source: str | None,
    questions: list[str],
) -> None:
    h = doc.add_heading(f"{number} {title}", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0E, 0x11, 0x16)

    if purpose:
        p = doc.add_paragraph()
        bold = p.add_run("Po co ta sekcja: ")
        bold.bold = True
        p.add_run(purpose)

    if quote and quote_source:
        p = doc.add_paragraph()
        bold = p.add_run("Co macie teraz na starej stronie:")
        bold.bold = True
        add_quote_box(doc, quote, quote_source)
    else:
        p = doc.add_paragraph()
        bold = p.add_run("Co macie teraz: ")
        bold.bold = True
        p.add_run("brak — ta sekcja nie ma odpowiednika na obecnej stronie. Klient wypełnia od zera.")

    p = doc.add_paragraph()
    bold = p.add_run("Pytania i pola do wypełnienia:")
    bold.bold = True

    for q in questions:
        bullet = doc.add_paragraph(q, style="List Bullet")
        for run in bullet.runs:
            run.font.size = Pt(11)
        add_answer_box(doc)


def build_doc() -> Document:
    doc = Document()

    # Marginesy 2.5 cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Domyślny font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_page_number_footer(doc)

    # Strona tytułowa
    title = doc.add_heading("Brief treściowy", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Nowa strona Vis Technologie", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    instr = doc.add_paragraph()
    instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = instr.add_run(
        "Wypełnij pola w obu częściach. Limity znaków w pytaniach są wskazówkami, "
        "nie regułami — lepiej napisać dłużej i dokładnie niż krótko byle co. "
        "Można odpowiadać w punktach."
    )
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.add_run("Data wypełnienia: ").bold = True
    add_answer_box(doc, "[data]")

    meta2 = doc.add_paragraph()
    meta2.add_run("Osoba wypełniająca: ").bold = True
    add_answer_box(doc, "[imię, nazwisko, rola]")

    add_page_break(doc)

    # Spis treści
    doc.add_heading("Spis treści", level=1)
    add_toc(doc)
    add_page_break(doc)

    # Część 1
    doc.add_heading(PART_1_TITLE, level=1)
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(PART_1_INTRO)
    intro_run.italic = True

    for num, title_, questions in PART_1:
        doc.add_heading(f"{num} {title_}", level=2)
        for q in questions:
            bullet = doc.add_paragraph(q, style="List Bullet")
            for run in bullet.runs:
                run.font.size = Pt(11)
            add_answer_box(doc)

    add_page_break(doc)

    # Część 2
    doc.add_heading(PART_2_TITLE, level=1)
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(PART_2_INTRO)
    intro_run.italic = True

    for num, title_, purpose, quote, quote_source, questions in PART_2:
        build_section(doc, num, title_, purpose, quote, quote_source, questions)

    add_page_break(doc)

    # Dopiski
    doc.add_heading(EPILOG_TITLE, level=1)
    for q in EPILOG:
        bullet = doc.add_paragraph(q, style="List Bullet")
        for run in bullet.runs:
            run.font.size = Pt(11)
        add_answer_box(doc)

    return doc


def main() -> int:
    repo_root = Path(__file__).resolve().parent.parent
    output = repo_root / "docs" / "brief-tresci-vistechnologie.docx"
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = build_doc()
    doc.save(str(output))

    size_kb = output.stat().st_size / 1024
    print(f"OK: {output} ({size_kb:.1f} KB)")
    print("Otwórz w Wordzie i naciśnij Ctrl+A, F9, żeby zaktualizować spis treści.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
